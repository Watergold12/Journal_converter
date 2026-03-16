# from docx import Document

# doc = Document("D:/volume_e_files/Projects/ECLearnix/journal_converter/data/Raw Copy/demo_title.docx")

# for para in doc.paragraphs:
#     for run in para.runs:
#         print("Text:", run.text)
#         print("Font:", run.font.name)
#         print("Size:", run.font.size)
#         print("Bold:", run.bold)
#         print("Italic:", run.italic)
#         print("----------------")

from docx import Document

doc = Document("D:/volume_e_files/Projects/ECLearnix/journal_converter/data/Raw Copy/demo_title.docx")

for para in doc.paragraphs:
    print(para.text)