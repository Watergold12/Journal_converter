# import re
# from docx import Document

# doc = Document("data/Raw Copy/2602320 Full Paper.docx")

# pattern = r'^\d+\.\d+'

# # sections that should never be treated as subheadings
# excluded_sections = {"REFERENCES", "ACKNOWLEDGEMENTS"}

# for para in doc.paragraphs:

#     text = para.text.strip()

#     if not text:
#         continue

#     # ---------------- TITLE DETECTION ----------------
#     is_title = False
#     for run in para.runs:
#         if run.font.size and run.font.size.pt > 12:
#             print("TITLE:", text)
#             is_title = True
#             break

#     if is_title:
#         continue

#     # ---------------- SECTION DETECTION ----------------
#     if text.isupper():
#         print("SECTION:", text)
#         continue

#     # ---------------- SUBHEADING DETECTION ----------------
#     if text.upper() in excluded_sections:
#         continue

#     numPr = para._element.xpath('./w:pPr/w:numPr')

#     if numPr:

#         # extract numbering level
#         level = numPr[0].xpath('./w:ilvl/@w:val')
#         numId = numPr[0].xpath('./w:numId/@w:val')

#         if level and int(level[0]) >= 1:

#             # try extracting the numbering prefix
#             number_match = re.match(pattern, text)

#             if number_match:
#                 number = number_match.group()
#                 heading_text = text[len(number):].strip()
#                 print(f"SUBHEADING: {number} {heading_text}")
#             else:
#                 print("SUBHEADING:", text)

#         continue

#     # ---------------- REGEX FALLBACK ----------------
#     if re.match(pattern, text):

#         if "reference" in text.lower():
#             continue

#         print("SUBHEADING:", text)

import re
from docx import Document

doc = Document("data/Raw Copy/2602320 Full Paper.docx")

pattern = r'^\d+\.\d+'
excluded_sections = {"REFERENCES", "ACKNOWLEDGEMENTS", "ACKNOWLEDGMENTS"}

# numbering counters
levels = [0] * 10

for para in doc.paragraphs:

    text = para.text.strip()

    if not text:
        continue

    # ---------------- TITLE DETECTION ----------------
    is_title = False
    for run in para.runs:
        if run.font.size and run.font.size.pt > 12:
            print("TITLE:", text)
            is_title = True
            break

    if is_title:
        continue

    # ---------------- SECTION DETECTION ----------------
    if text.isupper():

        # reset numbering when new section appears
        levels = [0] * 10

        print("SECTION:", text)
        continue

    if text.upper() in excluded_sections:
        continue

    # ---------------- XML NUMBERING DETECTION ----------------
    numPr = para._element.xpath('./w:pPr/w:numPr')

    if numPr:

        # Check if paragraph text is bold
        is_bold = any(run.bold for run in para.runs)

        if not is_bold:
            continue

        level = numPr[0].xpath('./w:ilvl/@w:val')

        if level:

            level = int(level[0])

            levels[level] += 1

            for i in range(level + 1, len(levels)):
                levels[i] = 0

            number = ".".join(str(levels[i]) for i in range(level + 1) if levels[i] > 0)

            print(f"SUBHEADING: {number} {text}")

        continue

    # ---------------- REGEX FALLBACK ----------------
    if re.match(pattern, text):

        if "reference" in text.lower():
            continue

        print("SUBHEADING:", text)