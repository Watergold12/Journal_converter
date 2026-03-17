from docx import Document

class DocumentParser:

    def __init__(self,file_path):

        self.file_path=file_path

        self.doc=Document(file_path)

    def get_paragraphs(self):

        return self.doc.paragraphs

    def get_document(self):

        return self.doc

    def get_tables(self):

        return self.doc.tables