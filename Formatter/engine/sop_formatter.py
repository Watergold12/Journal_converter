from io import BytesIO
import re
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class SOPFormatter:

    def __init__(self, source_doc):

        self.source = source_doc

        # NEW DOCUMENT (important)
        self.doc = Document()

        self.section_count = 0

        self.sub_count = 0

        self.current_section = 0

        self.numbering_enabled = True

        self.paragraph_image_rel_ids = set()

    def normalize_heading_key(self,text):

        text=re.sub(r"\s+"," ",text.strip())

        return text.lower().rstrip(":")


    def strip_number_prefix(self,text):

        return re.sub(r'^\d+(\.\d+)*[\.\)]?\s*','',text).strip()


    # GLOBAL FONT RULE
    def apply_font(self,para,size=12):

        for run in para.runs:

            run.font.name="Times New Roman"

            run.font.size=Pt(size)


    # TITLE RULE
    def add_title(self,text):

        para=self.doc.add_paragraph(text)

        self.apply_font(para,16)

        for run in para.runs:

            run.bold=True


    # AUTHOR RULE
    def add_author(self,text):

        para=self.doc.add_paragraph(text)

        self.apply_font(para,12)

        for run in para.runs:

            run.italic=True


    # SECTION RULE
    def add_heading(self,text):

        text=self.strip_number_prefix(text)

        heading_key=self.normalize_heading_key(text)

        NON_NUMBERED={"abstract","conclusion","acknowledgements","acknowledgments","references"}
        STOP_NUMBERING={"conclusion","acknowledgements","acknowledgments","references"}

        if heading_key in STOP_NUMBERING:

            self.numbering_enabled=False

            self.sub_count=0

        if heading_key in NON_NUMBERED or not self.numbering_enabled:

            para=self.doc.add_paragraph(text)

            self.apply_font(para)

            for run in para.runs:

                run.bold=True

                run.font.color.rgb=RGBColor(0,0,255)

            return


        self.section_count+=1

        self.sub_count=0

        self.current_section=self.section_count

        para=self.doc.add_paragraph(
            str(self.section_count)+". "+text
        )

        self.apply_font(para)

        for run in para.runs:

            run.bold=True

            run.font.color.rgb=RGBColor(0,0,255)


    # SUBSECTION RULE
    def add_subheading(self,text):

        text=self.strip_number_prefix(text)

        if self.current_section==0 or not self.numbering_enabled:

            para=self.doc.add_paragraph(text)

            self.apply_font(para)

            for run in para.runs:

                run.bold=True

            return

        self.sub_count+=1

        number=str(self.current_section)+"."+str(self.sub_count)

        para=self.doc.add_paragraph(
            number+" "+text
        )

        self.apply_font(para)

        for run in para.runs:

            run.bold=True


    # PARAGRAPH RULE
    def add_paragraph(self,text):

        para=self.doc.add_paragraph(text)

        self.apply_font(para)

        para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    # TABLE RULE
    def copy_tables(self):

        for table in self.source.tables:

            rows=len(table.rows)

            cols=len(table.columns)

            new_table=self.doc.add_table(rows,cols)

            # small tables stay in column
            if cols<4:

                pass

            # large tables full width
            else:

                section=self.doc.sections[0]

                cols_xml=section._sectPr.xpath('./w:cols')[0]

                cols_xml.set(qn('w:num'),'1')


            for i in range(rows):

                for j in range(cols):

                    text=table.cell(i,j).text

                    new_table.cell(i,j).text=text

                    for para in new_table.cell(i,j).paragraphs:

                        self.apply_font(para,10)


            # restore columns
            section=self.doc.sections[0]

            cols_xml=section._sectPr.xpath('./w:cols')[0]

            cols_xml.set(qn('w:num'),'2')


    # IMAGE RULE
    def add_image_by_rel_id(self,rel_id,width_emu=None):

        image_part=self.source.part.related_parts.get(rel_id)

        if image_part is None:

            return

        image_stream=BytesIO(image_part.blob)

        max_width=4000000

        if width_emu is not None:

            width_emu=min(int(width_emu),max_width)

            self.doc.add_picture(image_stream,width=Emu(width_emu))

        else:

            self.doc.add_picture(image_stream)


    def copy_images_from_paragraph(self,para):

        drawings=para._element.xpath('.//w:drawing')

        for drawing in drawings:

            rel_ids=drawing.xpath('.//a:blip/@r:embed')

            widths=drawing.xpath('.//wp:extent/@cx')

            width_emu=int(widths[0]) if widths else None

            for rel_id in rel_ids:

                self.paragraph_image_rel_ids.add(rel_id)

                self.add_image_by_rel_id(rel_id,width_emu)


    def copy_remaining_images(self):

        for shape in self.source.inline_shapes:

            rel_ids=shape._inline.xpath('.//a:blip/@r:embed')

            width_emu=int(shape.width) if shape.width is not None else None

            for rel_id in rel_ids:

                if rel_id in self.paragraph_image_rel_ids:

                    continue

                self.add_image_by_rel_id(rel_id,width_emu)



    # COLUMN RULE
    def apply_double_column(self):

        section=self.doc.sections[0]

        cols=section._sectPr.xpath('./w:cols')[0]

        cols.set(qn('w:num'),'2')


    # MAIN ENGINE
    def build(self,detector):

        for para in self.source.paragraphs:

            self.copy_images_from_paragraph(para)

            text=para.text.strip()

            if text=="":

                continue


            block=detector.classify(para)


            if block=="TITLE":

                self.add_title(text)


            elif block=="AUTHOR":

                self.add_author(text)


            elif block=="HEADING":

                self.add_heading(text)


            elif block=="SUBHEADING":

                self.add_subheading(text)


            else:

                self.add_paragraph(text)

        # Handles inline shapes that may not be discoverable via paragraphs.
        self.copy_remaining_images()

        self.copy_tables()

        self.apply_double_column()


        return self.doc
