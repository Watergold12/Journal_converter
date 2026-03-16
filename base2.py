import pdfplumber
from docx import Document

doc = Document()

def extract_text_from_pdf(pdf_path):
    all_text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()

            if text:
                doc.add_paragraph(f"--- Page {i} ---")
                doc.add_paragraph(text)
        doc.save("extracted_text.docx")
    return all_text

file_path = "ZEPHYR_VISHAL_AA_24AD124_REPORT.pdf" 

extract_text_from_pdf(file_path)
