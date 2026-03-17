# from docx import Document

# doc = Document("data/Raw_Copy/2602320_Full_Paper.docx")

# for para in doc.paragraphs:
#     print(para.text)

from docx import Document

doc = Document("data/Raw_Copy/2602320_Full_Paper.docx")

for para in doc.paragraphs:

    text = para.text.strip()

    if not text:
        continue

    # detect formatting
    is_bold = any(run.bold for run in para.runs)

    font_sizes = [run.font.size.pt for run in para.runs if run.font.size]
    max_size = max(font_sizes) if font_sizes else None

    colors = [run.font.color.rgb for run in para.runs if run.font.color.rgb]

    # -------- TITLE --------
    if max_size and max_size > 12:
        print("TITLE:", text)
        continue

    # -------- SECTION --------
    if is_bold and text.isupper():
        print("SECTION:", text)
        continue

    # -------- SUBHEADING --------
    if is_bold and not text.isupper():
        print("SUBHEADING:", text)
        continue

    # -------- CONTENT --------
    print("CONTENT:", text)